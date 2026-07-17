"""
main.py
-------
CLI entry point for the PII Redaction Tool.

Usage:
    python main.py input.docx output.docx
    python main.py input.docx output.docx --ground-truth reports/ground_truth.json
    python main.py input.docx output.docx --report-dir reports/

The tool:
    1. Loads input.docx
    2. Runs the hybrid regex + spaCy NER detection pipeline over every
       paragraph, table cell, header, and footer
    3. Replaces each detected PII value with a consistent fake value
    4. Saves the result to output.docx
    5. Writes a redaction summary (reports/redaction_summary.md) and, if a
       ground-truth file is supplied, precision/recall/F1 evaluation reports
"""

from __future__ import annotations

import argparse
import json
import sys
import time
from pathlib import Path

from detector import PIIDetector
from evaluator import Evaluator
from fake_generator import FakeGenerator
from redactor import DocxRedactor
from utils import setup_logger

logger = setup_logger("main")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Redact PII from a .docx document, replacing it with realistic fake values."
    )
    parser.add_argument("input", help="Path to the source .docx file")
    parser.add_argument("output", help="Path to write the redacted .docx file")
    parser.add_argument("--ground-truth", default=None,
                         help="Optional path to a JSON ground-truth annotation file "
                              "for precision/recall/F1 evaluation")
    parser.add_argument("--report-dir", default="reports",
                         help="Directory to write summary/evaluation reports into "
                              "(default: reports/)")
    parser.add_argument("--mapping-log", default=None,
                         help="Optional path to write the full original->fake value "
                              "mapping as JSON (useful for auditing)")
    return parser.parse_args()


def main() -> int:
    args = parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)
    report_dir = Path(args.report_dir)
    report_dir.mkdir(parents=True, exist_ok=True)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if not input_path.exists():
        logger.error("Input file not found: %s", input_path)
        return 1

    logger.info("Initializing detection pipeline (spaCy + regex)...")
    t0 = time.time()
    detector = PIIDetector()
    fake_generator = FakeGenerator()
    redactor = DocxRedactor(detector=detector, fake_generator=fake_generator)
    logger.info("Pipeline ready in %.1fs", time.time() - t0)

    t0 = time.time()
    stats = redactor.redact_file(str(input_path), str(output_path))
    elapsed = time.time() - t0
    logger.info("Redaction complete in %.1fs", elapsed)

    # ---- Redaction summary report ----
    summary_path = report_dir / "redaction_summary.md"
    with open(summary_path, "w", encoding="utf-8") as f:
        f.write("# PII Redaction — Run Summary\n\n")
        f.write(f"- Input: `{input_path}`\n")
        f.write(f"- Output: `{output_path}`\n")
        f.write(f"- Processing time: {elapsed:.1f}s\n")
        f.write(f"- Total entities redacted: **{stats.total_entities}**\n")
        f.write(f"- Unique original->fake mappings: **{len(fake_generator.mapping)}**\n\n")
        f.write("| Entity Type | Count |\n|---|---|\n")
        for label in sorted(stats.by_label):
            f.write(f"| {label} | {stats.by_label[label]} |\n")
    logger.info("Wrote redaction summary: %s", summary_path)

    if args.mapping_log:
        with open(args.mapping_log, "w", encoding="utf-8") as f:
            json.dump(
                [{"original": o, "fake": fk, "label": lb} for o, fk, lb in redactor.mapping_log],
                f, indent=2, ensure_ascii=False,
            )
        logger.info("Wrote mapping log: %s", args.mapping_log)

    # ---- Optional evaluation against ground truth ----
    if args.ground_truth:
        gt_path = Path(args.ground_truth)
        if not gt_path.exists():
            logger.warning("Ground truth file not found, skipping evaluation: %s", gt_path)
        else:
            logger.info("Running evaluation against ground truth: %s", gt_path)
            evaluator = Evaluator()
            ground_truth = evaluator.load_ground_truth(str(gt_path))

            # Re-detect on the raw extracted text of the *original* document
            # so evaluation reflects detector performance independent of
            # docx run-splitting quirks.
            from docx import Document

            def table_text(table):
                chunks = []
                for row in table.rows:
                    for cell in row.cells:
                        chunks.append("\n".join(p.text for p in cell.paragraphs))
                        for nested in cell.tables:
                            chunks.append(table_text(nested))
                return "\n".join(chunks)

            doc = Document(str(input_path))
            text_chunks = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                text_chunks.append(table_text(table))
            for section in doc.sections:
                for part in (section.header, section.footer,
                             section.first_page_header, section.first_page_footer,
                             section.even_page_header, section.even_page_footer):
                    if part is not None:
                        text_chunks.extend(p.text for p in part.paragraphs)
                        for table in part.tables:
                            text_chunks.append(table_text(table))
            full_text = "\n".join(text_chunks)

            predicted = detector.detect(full_text)
            evaluator.evaluate(predicted, ground_truth)
            evaluator.write_reports(
                str(report_dir / "evaluation_report.csv"),
                str(report_dir / "evaluation_report.md"),
            )
            logger.info("Overall accuracy: %.3f", evaluator.overall_accuracy())

    logger.info("Done.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
