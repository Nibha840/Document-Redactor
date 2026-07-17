"""
redactor.py
-----------
Reads a source .docx, detects PII, and writes a new .docx with every
detected PII value replaced by a consistent fake value — while preserving
the original document's formatting (fonts, bold/italic, styles, tables,
headers/footers) as much as possible.

Approach:
    python-docx exposes text at the *run* level (`paragraph.runs`), and Word
    frequently splits a single visible sentence across many runs (due to
    spell-check boundaries, revision marks, etc.). To detect entities
    reliably we first flatten each paragraph to its full plain text, run
    detection on that flattened string, then map the resulting character
    offsets back onto the underlying runs and rewrite run.text in place.
    This preserves per-run formatting because we never delete/recreate
    runs — we only mutate their `.text`.

Coverage:
    - Body paragraphs
    - Table cells (including nested tables)
    - Headers / footers (all sections)
    - Text boxes are embedded as nested XML and are intentionally out of
      scope for python-docx run-level editing; see README "Tradeoffs".
"""

from __future__ import annotations

from dataclasses import dataclass
from typing import List

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph

from detector import PIIDetector
from fake_generator import FakeGenerator
from utils import Entity, setup_logger

logger = setup_logger(__name__)


@dataclass
class RedactionStats:
    total_entities: int = 0
    by_label: dict = None

    def __post_init__(self):
        if self.by_label is None:
            self.by_label = {}

    def record(self, entities: List[Entity]) -> None:
        self.total_entities += len(entities)
        for e in entities:
            self.by_label[e.label] = self.by_label.get(e.label, 0) + 1


class DocxRedactor:
    """Orchestrates PII detection + in-place redaction of a .docx file."""

    def __init__(self, detector: PIIDetector | None = None,
                 fake_generator: FakeGenerator | None = None) -> None:
        self.detector = detector or PIIDetector()
        self.fake_generator = fake_generator or FakeGenerator()
        self.stats = RedactionStats()
        # collected (original -> fake, label) tuples for the evaluation /
        # mapping report
        self.mapping_log: List[tuple] = []

    # ------------------------------------------------------------------ #
    # Public API
    # ------------------------------------------------------------------ #
    def redact_file(self, input_path: str, output_path: str) -> RedactionStats:
        logger.info("Loading document: %s", input_path)
        doc = Document(input_path)

        logger.info("Redacting body paragraphs...")
        self._redact_paragraphs(doc.paragraphs)

        logger.info("Redacting tables...")
        for table in doc.tables:
            self._redact_table(table)

        logger.info("Redacting headers and footers...")
        for section in doc.sections:
            for part in (section.header, section.footer,
                         section.first_page_header, section.first_page_footer,
                         section.even_page_header, section.even_page_footer):
                if part is not None:
                    self._redact_paragraphs(part.paragraphs)
                    for table in part.tables:
                        self._redact_table(table)

        doc.save(output_path)
        logger.info("Saved redacted document: %s", output_path)
        logger.info("Total entities redacted: %d | breakdown: %s",
                     self.stats.total_entities, self.stats.by_label)
        return self.stats

    # ------------------------------------------------------------------ #
    # Structural traversal
    # ------------------------------------------------------------------ #
    def _redact_table(self, table: Table) -> None:
        for row in table.rows:
            for cell in row.cells:
                self._redact_paragraphs(cell.paragraphs)
                for nested_table in cell.tables:
                    self._redact_table(nested_table)

    def _redact_paragraphs(self, paragraphs: List[Paragraph]) -> None:
        for paragraph in paragraphs:
            self._redact_paragraph(paragraph)

    # ------------------------------------------------------------------ #
    # Core run-level redaction
    # ------------------------------------------------------------------ #
    def _redact_paragraph(self, paragraph: Paragraph) -> None:
        runs = paragraph.runs
        if not runs:
            return

        full_text = "".join(r.text for r in runs)
        if not full_text.strip():
            return

        entities = self.detector.detect(full_text)
        if not entities:
            return

        # Build run boundary offsets so we can map character spans -> runs
        run_spans = []
        cursor = 0
        for r in runs:
            run_spans.append((cursor, cursor + len(r.text)))
            cursor += len(r.text)

        # Apply replacements back-to-front to keep earlier offsets valid.
        for entity in sorted(entities, key=lambda e: e.start, reverse=True):
            fake_value = self.fake_generator.get_fake(entity.text, entity.label)
            self.mapping_log.append((entity.text, fake_value, entity.label))
            self._apply_replacement(runs, run_spans, entity, fake_value)

        self.stats.record(entities)

    @staticmethod
    def _apply_replacement(runs, run_spans, entity: Entity, fake_value: str) -> None:
        """Replace the character span [entity.start, entity.end) — which may
        cross multiple runs — with `fake_value`, placing all replacement
        text in the first affected run and blanking the remainder so
        formatting of the first run is preserved."""
        start, end = entity.start, entity.end
        affected = [
            i for i, (s, e) in enumerate(run_spans)
            if s < end and e > start
        ]
        if not affected:
            return

        first_idx = affected[0]
        for idx in affected:
            r_start, r_end = run_spans[idx]
            local_start = max(start, r_start) - r_start
            local_end = min(end, r_end) - r_start
            original_run_text = runs[idx].text
            if idx == first_idx:
                new_text = (
                    original_run_text[:local_start] + fake_value + original_run_text[local_end:]
                )
            else:
                new_text = original_run_text[:local_start] + original_run_text[local_end:]
            runs[idx].text = new_text

        # Because run lengths changed but run_spans is a snapshot for this
        # paragraph only (recomputed fresh per paragraph call), and we
        # process entities back-to-front within the paragraph, offsets for
        # not-yet-processed (earlier) entities in this same paragraph
        # remain valid.
